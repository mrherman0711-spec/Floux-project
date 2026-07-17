#!/usr/bin/env python3
"""Convert the Uber report HTML into a formatted .docx using python-docx."""
import re
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "/Users/maxherman/Desktop/Floux Project/uber-org-psych-report.html"
OUT = "/Users/maxherman/Desktop/Floux Project/uber-org-psych-report.docx"


class Node:
    def __init__(self, tag, attrs=None):
        self.tag = tag
        self.attrs = dict(attrs or [])
        self.children = []  # list of Node or str


class TreeBuilder(HTMLParser):
    VOID = {"br", "hr", "meta"}

    def __init__(self):
        super().__init__()
        self.root = Node("root")
        self.stack = [self.root]

    def handle_starttag(self, tag, attrs):
        if tag in self.VOID:
            self.stack[-1].children.append(Node(tag, attrs))
            return
        node = Node(tag, attrs)
        self.stack[-1].children.append(node)
        self.stack.append(node)

    def handle_startendtag(self, tag, attrs):
        self.stack[-1].children.append(Node(tag, attrs))

    def handle_endtag(self, tag):
        if tag in self.VOID:
            return
        for i in range(len(self.stack) - 1, 0, -1):
            if self.stack[i].tag == tag:
                del self.stack[i:]
                break

    def handle_data(self, data):
        if data.strip("\n\t ") == "" and "\n" in data:
            return
        self.stack[-1].children.append(data)


def collect_inline(node):
    """Yield (text, bold, italic) runs from inline content."""
    runs = []

    def walk(n, bold, italic):
        if isinstance(n, str):
            text = re.sub(r"\s+", " ", n)
            if text:
                runs.append((text, bold, italic))
            return
        b = bold or n.tag in ("strong", "b")
        i = italic or n.tag in ("em", "i")
        if n.tag == "br":
            runs.append(("\n", bold, italic))
            return
        for c in n.children:
            walk(c, b, i)

    for c in node.children:
        walk(c, False, False)
    return runs


def add_inline_paragraph(doc, node, style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, bold, italic in collect_inline(node):
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    return p


def render(doc, node):
    for child in node.children:
        if isinstance(child, str):
            continue
        tag = child.tag
        if tag == "h1":
            doc.add_heading(collect_text(child), level=0)
        elif tag == "h2":
            doc.add_heading(collect_text(child), level=1)
        elif tag == "h3":
            doc.add_heading(collect_text(child), level=2)
        elif tag == "p":
            align = "center" if "center" in child.attrs.get("style", "") else None
            add_inline_paragraph(doc, child, align=align)
        elif tag == "ul":
            for li in child.children:
                if isinstance(li, Node) and li.tag == "li":
                    add_inline_paragraph(doc, li, style="List Bullet")
        elif tag == "ol":
            for li in child.children:
                if isinstance(li, Node) and li.tag == "li":
                    add_inline_paragraph(doc, li, style="List Number")
        elif tag == "table":
            render_table(doc, child)
        elif tag == "hr":
            doc.add_page_break()
        elif tag in ("body", "html", "head", "div"):
            render(doc, child)


def collect_text(node):
    return "".join(t for t, _, _ in collect_inline(node)).strip()


def render_table(doc, table_node):
    rows = [r for r in table_node.children if isinstance(r, Node) and r.tag == "tr"]
    if not rows:
        return
    ncols = max(
        len([c for c in r.children if isinstance(c, Node) and c.tag in ("td", "th")])
        for r in rows
    )
    t = doc.add_table(rows=0, cols=ncols)
    t.style = "Light Grid Accent 1"
    for r in rows:
        cells = [c for c in r.children if isinstance(c, Node) and c.tag in ("td", "th")]
        row_cells = t.add_row().cells
        for idx, cell in enumerate(cells):
            if idx >= ncols:
                break
            tc = row_cells[idx]
            tc.paragraphs[0].text = ""
            runs = collect_inline(cell)
            para = tc.paragraphs[0]
            for text, bold, italic in runs:
                run = para.add_run(text)
                run.bold = bold or cell.tag == "th"
                run.italic = italic


def main():
    with open(SRC, encoding="utf-8") as f:
        html = f.read()
    builder = TreeBuilder()
    builder.feed(html)
    doc = Document()
    # base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    render(doc, builder.root)
    doc.save(OUT)
    print("Saved", OUT)


if __name__ == "__main__":
    main()
